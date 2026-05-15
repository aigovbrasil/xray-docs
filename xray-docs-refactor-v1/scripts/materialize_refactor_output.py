#!/usr/bin/env python3
"""
Materializa uma saída de refatoração documental a partir de um work_order.json.

Este script é desenhado para ser usado como arquivo de apoio de um Skill Anthropic.
Ele NÃO chama nenhum provedor LLM. Toda decisão semântica deve vir no work_order.

Entrada:
- work_order.json com a lista de documentos já classificados/refatorados.
- config JSON opcional com phase_map, naming e extensões permitidas.

Saída:
- árvore por fase
- documentos refatorados (.md, .json, .docx)
- *.meta.json por documento
- manifest.json
- report.json
- output_refatorado.zip
"""

from __future__ import annotations

import argparse
import dataclasses
import json
import re
import shutil
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

try:
    import docx  # python-docx
except Exception:
    docx = None

DEFAULT_CONFIG: Dict[str, Any] = {
    "project_name": "negocio",
    "output_root": "NEGOCIO",
    "write_original_copy": False,
    "naming": {
        "folder_len_min": 12,
        "folder_len_max": 30,
        "filename_len_min": 25,
        "filename_len_max": 48,
        "max_path_length": 180,
        "allowed_extensions": [".md", ".json", ".docx"],
    },
    "phase_map": {
        "00_admin": "00_admin",
        "01_estrategia": "01_estrategia",
        "02_discovery_validacao": "02_discovery-validacao",
        "03_mvp_produto": "03_mvp-produto",
        "04_vendas_marketing": "04_vendas-marketing",
        "05_operacao_entrega": "05_operacao-entrega",
        "06_financeiro_legal": "06_financeiro-legal",
        "07_growth_analytics": "07_growth-analytics",
        "99_arquivo": "99_arquivo",
    },
    "doc_types": [
        "readme-operacional",
        "indice-mestre",
        "stack-acessos",
        "lean-canvas",
        "hipoteses-criticas",
        "icp-card",
        "metricas-norte",
        "roteiro-entrevista",
        "log-entrevistas",
        "experimentos",
        "oferta-onepager",
        "pipeline-validacao",
        "prd-lite",
        "backlog-mvp",
        "fluxos-ux",
        "dicionario-de-dados",
        "mensagem-comercial-base",
        "faq-objecoes",
        "crm-comercial",
        "dashboard-comercial",
        "sop-onboarding",
        "sop-entrega",
        "sop-suporte",
        "operacoes-semana",
        "change-log",
        "fluxo-caixa",
        "dre-gerencial",
        "precificacao",
        "roadmap-now-next-later",
        "experimentos-growth",
        "dashboard-negocio",
        "outro",
    ],
    "target_ext_by_doc_type": {
        "default": ".md",
        "stack-acessos": ".json",
        "log-entrevistas": ".json",
        "experimentos": ".json",
        "pipeline-validacao": ".json",
        "backlog-mvp": ".json",
        "dicionario-de-dados": ".json",
        "crm-comercial": ".json",
        "dashboard-comercial": ".json",
        "operacoes-semana": ".json",
        "fluxo-caixa": ".json",
        "dre-gerencial": ".json",
        "precificacao": ".json",
        "experimentos-growth": ".json",
        "dashboard-negocio": ".json",
    },
}


@dataclasses.dataclass
class FileResult:
    source_path: str
    detected_ext: str
    phase: str
    doc_type: str
    title: str
    output_rel_path: str
    status: str
    warnings: List[str]
    output_chars: int
    confidence: float


def deep_merge(base: Dict[str, Any], new: Dict[str, Any]) -> None:
    for key, value in new.items():
        if isinstance(value, dict) and isinstance(base.get(key), dict):
            deep_merge(base[key], value)
        else:
            base[key] = value


def load_json(path: Path) -> Dict[str, Any]:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def load_config(path: Optional[str]) -> Dict[str, Any]:
    cfg = json.loads(json.dumps(DEFAULT_CONFIG))
    if path:
        deep_merge(cfg, load_json(Path(path)))
    return cfg


def slugify(text: str, keep_underscore: bool = False) -> str:
    text = str(text).lower().strip()
    text = text.replace("ç", "c")
    text = re.sub(r"[áàãâä]", "a", text)
    text = re.sub(r"[éèêë]", "e", text)
    text = re.sub(r"[íìîï]", "i", text)
    text = re.sub(r"[óòõôö]", "o", text)
    text = re.sub(r"[úùûü]", "u", text)
    allowed = r"[^a-z0-9_\-]+" if keep_underscore else r"[^a-z0-9\-]+"
    text = re.sub(allowed, "-", text)
    text = re.sub(r"-+", "-", text)
    text = re.sub(r"_+", "_", text)
    return text.strip("-_")


def trim_with_budget(base: str, max_len: int, preserve_suffix: str) -> str:
    """Garante que o sufixo seja preservado ao cortar o nome."""
    if len(base) <= max_len:
        return base
    budget = max_len - len(preserve_suffix)
    if budget < 8:
        budget = max_len
        preserve_suffix = ""
    head = base[:budget].rstrip("-_")
    return (head + preserve_suffix).rstrip("-_")


def ensure_len(text: str, min_len: int, max_len: int, filler: str, preserve_suffix: str = "") -> str:
    text = trim_with_budget(text, max_len=max_len, preserve_suffix=preserve_suffix)
    if len(text) >= min_len:
        return text
    filler = slugify(filler) or "doc"
    while len(text) < min_len:
        candidate = (text + "-" + filler).rstrip("-_")
        text = trim_with_budget(candidate, max_len=max_len, preserve_suffix=preserve_suffix)
        if len(text) == max_len or filler == "doc":
            break
        filler = "doc"
    return text[:max_len].rstrip("-_")


def render_json_document(payload: Dict[str, Any]) -> str:
    content_value: Any = payload["refactored_content"]
    if isinstance(content_value, str):
        try:
            content_value = json.loads(content_value)
        except Exception:
            pass
    body = {
        "title": payload["title"],
        "objective": payload.get("objective", ""),
        "doc_type": payload["doc_type"],
        "phase_key": payload["phase_key"],
        "confidence": payload.get("confidence", 0.0),
        "content": content_value,
        "warnings": payload.get("warnings", []),
    }
    return json.dumps(body, ensure_ascii=False, indent=2)


def write_docx(path: Path, payload: Dict[str, Any]) -> int:
    if docx is None:
        raise RuntimeError("python-docx não instalado. Rode: pip install python-docx")
    document = docx.Document()
    document.add_heading(payload["title"], level=1)
    if payload.get("objective"):
        document.add_paragraph(payload["objective"])
    document.add_paragraph("")
    for raw_line in str(payload["refactored_content"]).splitlines():
        line = raw_line.rstrip()
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return len(str(payload["refactored_content"]))


def write_output_document(path: Path, payload: Dict[str, Any]) -> int:
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.suffix == ".json":
        rendered = render_json_document(payload)
        path.write_text(rendered, encoding="utf-8")
        return len(rendered)
    if path.suffix == ".docx":
        return write_docx(path, payload)
    content = str(payload["refactored_content"])
    path.write_text(content, encoding="utf-8")
    return len(content)


def zip_directory(source_dir: Path, output_zip: Path) -> None:
    with zipfile.ZipFile(output_zip, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for file_path in source_dir.rglob("*"):
            if file_path.is_file() and file_path != output_zip:
                zf.write(file_path, file_path.relative_to(source_dir))


def validate_document(document: Dict[str, Any], cfg: Dict[str, Any]) -> Tuple[Dict[str, Any], List[str]]:
    warnings: List[str] = list(document.get("warnings") or [])

    source_path = str(document.get("source_path") or "").strip()
    detected_ext = str(document.get("detected_ext") or "").strip() or Path(source_path).suffix.lower()

    phase_key = str(document.get("phase_key") or "99_arquivo").strip()
    if phase_key not in cfg["phase_map"]:
        warnings.append(f"phase_key inválido: {phase_key}. Ajustado para 99_arquivo.")
        phase_key = "99_arquivo"

    doc_type = slugify(document.get("doc_type") or "outro")
    if doc_type not in cfg["doc_types"]:
        warnings.append(f"doc_type inválido: {doc_type}. Ajustado para outro.")
        doc_type = "outro"

    title = re.sub(r"\s+", " ", str(document.get("title") or doc_type)).strip()
    if not title:
        title = doc_type
        warnings.append("Título vazio; preenchido a partir de doc_type.")

    objective = re.sub(r"\s+", " ", str(document.get("objective") or "")).strip()
    target_extension = str(document.get("target_extension") or cfg["target_ext_by_doc_type"].get(doc_type, ".md")).strip().lower()
    if target_extension not in cfg["naming"]["allowed_extensions"]:
        warnings.append(f"Extensão alvo inválida: {target_extension}. Ajustada para .md.")
        target_extension = ".md"

    refactored_content = document.get("refactored_content")
    if refactored_content is None or (isinstance(refactored_content, str) and not refactored_content.strip()):
        warnings.append("Conteúdo refatorado vazio. Inserido placeholder INCOMPLETO.")
        refactored_content = "# INCOMPLETO\n\nConteúdo insuficiente para materialização."

    try:
        confidence = float(document.get("confidence", 0.0))
    except Exception:
        confidence = 0.0
        warnings.append("Confidence inválido. Ajustado para 0.0.")
    confidence = max(0.0, min(1.0, confidence))
    if confidence < 0.60:
        warnings.append("Confidence abaixo de 0.60; revisar manualmente.")

    normalized = {
        "source_path": source_path,
        "detected_ext": detected_ext,
        "phase_key": phase_key,
        "doc_type": doc_type,
        "title": title,
        "objective": objective,
        "target_extension": target_extension,
        "refactored_content": refactored_content,
        "warnings": warnings,
        "confidence": confidence,
    }
    return normalized, warnings


def make_output_filename(payload: Dict[str, Any], cfg: Dict[str, Any]) -> str:
    phase_key = payload["phase_key"]
    doc_type = payload["doc_type"]
    title = payload["title"]
    ext = payload["target_extension"]
    date_part = datetime.now().strftime("%Y-%m")
    version_suffix = "-v01"
    base = f"{phase_key[:2]}_{slugify(doc_type)}_{date_part}_{slugify(title)}"
    base = ensure_len(
        text=base,
        min_len=int(cfg["naming"]["filename_len_min"]),
        max_len=int(cfg["naming"]["filename_len_max"]),
        filler=doc_type,
        preserve_suffix=version_suffix,
    )
    if not base.endswith(version_suffix):
        base = trim_with_budget(base, int(cfg["naming"]["filename_len_max"]), version_suffix)
        if not base.endswith(version_suffix):
            base = f"{base}{version_suffix}".rstrip("-_")
            base = trim_with_budget(base, int(cfg["naming"]["filename_len_max"]), "")
    return base + ext


def create_scaffold(out_root: Path, cfg: Dict[str, Any]) -> None:
    for folder in cfg["phase_map"].values():
        (out_root / folder).mkdir(parents=True, exist_ok=True)
    readme = out_root / "00_admin" / f"00_readme-operacional_{datetime.now().strftime('%Y-%m')}_estrutura-gerada_v01.md"
    if not readme.exists():
        readme.write_text(
            "# README Operacional\n\n"
            "Estrutura gerada por skill Anthropic.\n"
            "Cada documento possui um arquivo `.meta.json` correspondente.\n",
            encoding="utf-8",
        )


def process_documents(documents: Iterable[Dict[str, Any]], cfg: Dict[str, Any], out_root: Path) -> List[FileResult]:
    results: List[FileResult] = []

    for idx, raw_document in enumerate(documents, start=1):
        payload, warnings = validate_document(raw_document, cfg)
        phase_folder = cfg["phase_map"][payload["phase_key"]]
        filename = make_output_filename(payload, cfg)
        output_path = out_root / phase_folder / filename

        if len(str(output_path)) > int(cfg["naming"]["max_path_length"]):
            warnings.append("Caminho acima do limite recomendado; título reduzido automaticamente.")
            payload["title"] = payload["title"][:18].rstrip()
            filename = make_output_filename(payload, cfg)
            output_path = out_root / phase_folder / filename

        # Resolve filename collision: append _002, _003 ... if path already exists
        if output_path.exists():
            stem = output_path.stem
            suffix = output_path.suffix
            counter = 2
            while output_path.exists():
                new_stem = re.sub(r"-v\d+$", f"-v{counter:02d}", stem)
                if new_stem == stem:
                    new_stem = f"{stem}-v{counter:02d}"
                output_path = output_path.with_name(new_stem + suffix)
                counter += 1

        output_chars = write_output_document(output_path, payload)

        if cfg.get("write_original_copy") and payload["source_path"]:
            source_path = Path(payload["source_path"])
            if source_path.exists() and source_path.is_file():
                originals_dir = out_root / "99_arquivo" / "originais"
                originals_dir.mkdir(parents=True, exist_ok=True)
                shutil.copy2(source_path, originals_dir / source_path.name)
            else:
                warnings.append("source_path não encontrado; cópia do original foi ignorada.")

        meta_payload = {
            "source_path": payload["source_path"],
            "source_extension": payload["detected_ext"],
            "phase_key": payload["phase_key"],
            "phase_folder": phase_folder,
            "doc_type": payload["doc_type"],
            "title": payload["title"],
            "objective": payload["objective"],
            "confidence": payload["confidence"],
            "warnings": warnings,
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "sequence": idx,
        }
        meta_path = output_path.with_suffix(output_path.suffix + ".meta.json")
        meta_path.write_text(json.dumps(meta_payload, ensure_ascii=False, indent=2), encoding="utf-8")

        results.append(
            FileResult(
                source_path=payload["source_path"],
                detected_ext=payload["detected_ext"],
                phase=payload["phase_key"],
                doc_type=payload["doc_type"],
                title=payload["title"],
                output_rel_path=str(output_path.relative_to(out_root)),
                status="ok",
                warnings=warnings,
                output_chars=output_chars,
                confidence=payload["confidence"],
            )
        )

    return results


def write_manifest(out_root: Path, results: List[FileResult], cfg: Dict[str, Any], work_order_path: Path) -> None:
    manifest = {
        "project_name": cfg["project_name"],
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "work_order_path": str(work_order_path),
        "total_files": len(results),
        "results": [dataclasses.asdict(result) for result in results],
    }
    (out_root / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")


def write_report(out_root: Path, results: List[FileResult]) -> None:
    by_phase: Dict[str, int] = {}
    by_doc_type: Dict[str, int] = {}
    warnings_total = 0

    for result in results:
        by_phase[result.phase] = by_phase.get(result.phase, 0) + 1
        by_doc_type[result.doc_type] = by_doc_type.get(result.doc_type, 0) + 1
        warnings_total += len(result.warnings)

    report = {
        "summary": {
            "files_processed": len(results),
            "warnings": warnings_total,
        },
        "by_phase": by_phase,
        "by_doc_type": by_doc_type,
    }
    (out_root / "report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Materializa documentos refatorados a partir de um work_order.json")
    parser.add_argument("--work-order", required=True, help="Caminho para o work_order.json")
    parser.add_argument("--output-dir", required=True, help="Pasta de saída")
    parser.add_argument("--config", help="Arquivo JSON opcional de configuração")
    parser.add_argument("--clean-output", action="store_true", help="Remove a pasta de saída antes de gerar")
    parser.add_argument("--validate-only", action="store_true", help="Somente valida o work_order e sai")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    cfg = load_config(args.config)

    work_order_path = Path(args.work_order).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()

    if not work_order_path.exists() or not work_order_path.is_file():
        print(f"ERRO: work_order inválido: {work_order_path}", file=sys.stderr)
        return 1

    work_order = load_json(work_order_path)
    documents = work_order.get("documents")
    if not isinstance(documents, list) or not documents:
        print("ERRO: work_order.json precisa conter um array não vazio em 'documents'.", file=sys.stderr)
        return 2

    deep_merge(cfg, {k: v for k, v in work_order.items() if k != "documents"})

    if args.validate_only:
        local_warnings = 0
        for idx, document in enumerate(documents, start=1):
            _, warnings = validate_document(document, cfg)
            local_warnings += len(warnings)
            if warnings:
                src = document.get("source_path", f"doc #{idx}")
                print(f"  [{idx}] {src}")
                for w in warnings:
                    print(f"      ⚠  {w}")
        print(f"\nValidação concluída. Documentos: {len(documents)} | Warnings: {local_warnings}")
        return 0

    if args.clean_output and output_dir.exists():
        shutil.rmtree(output_dir)

    out_root = output_dir / cfg["output_root"]
    out_root.mkdir(parents=True, exist_ok=True)
    create_scaffold(out_root, cfg)

    results = process_documents(documents, cfg, out_root)
    write_manifest(out_root, results, cfg, work_order_path)
    write_report(out_root, results)

    zip_path = output_dir / "output_refatorado.zip"
    zip_directory(out_root, zip_path)

    print("Processamento concluído.")
    print(f"Arquivos processados: {len(results)}")
    print(f"ZIP final: {zip_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
