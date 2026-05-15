#!/usr/bin/env python3
"""
build_corpus.py — Business Document Pipeline Output Packager

Assembles all approved documents into:
  1. A formatted workbook (.docx)
  2. An organized folder structure (.zip)

Usage:
    python scripts/build_corpus.py \
        --corpus-dir ./output/corpus \
        --project-name "My Project" \
        --profile product \
        --output-dir ./output/final

Dependencies:
    pip install python-docx
"""

import argparse
import json
import os
import re
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Run: pip install python-docx")


# ─── Phase ordering ───────────────────────────────────────────────────────────

PHASE_ORDER = {
    "discovery": ["vision-framing", "mrd-lite", "prfaq-lite", "business-case", "brd-lite"],
    "product":   ["prd-lite", "frd-lite", "nfr-onepager", "adr-decision-log"],
    "execution": ["roadmap", "user-stories", "backlog", "release-plan"],
    "operations":["sop", "runbook", "data-integration-spec"],
}

DOCUMENT_TITLES = {
    "vision-framing":       "Vision Framing",
    "mrd-lite":             "Market Requirements Document",
    "prfaq-lite":           "PR/FAQ",
    "business-case":        "Business Case",
    "brd-lite":             "Business Requirements Document",
    "prd-lite":             "Product Requirements Document",
    "frd-lite":             "Functional Requirements Document",
    "nfr-onepager":         "Non-Functional Requirements",
    "adr-decision-log":     "Architecture Decision Records",
    "roadmap":              "Product Roadmap",
    "user-stories":         "User Stories",
    "backlog":              "Product Backlog",
    "release-plan":         "Release Plan",
    "sop":                  "Standard Operating Procedure",
    "runbook":              "Runbook",
    "data-integration-spec":"Data Integration Specification",
}


# ─── Helpers ──────────────────────────────────────────────────────────────────

def find_corpus_files(corpus_dir: Path) -> dict:
    """Scan corpus_dir for approved markdown documents."""
    docs = {}
    for phase, doc_types in PHASE_ORDER.items():
        for doc_type in doc_types:
            candidates = list(corpus_dir.glob(f"**/{doc_type}*.md"))
            candidates += list(corpus_dir.glob(f"**/{doc_type}.txt"))
            if candidates:
                docs[doc_type] = candidates[0]
    return docs


def read_document(path: Path) -> str:
    """Read a document file and return its content."""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def strip_frontmatter(content: str) -> str:
    """Remove YAML frontmatter if present."""
    if content.startswith("---"):
        parts = content.split("---", 2)
        if len(parts) >= 3:
            return parts[2].strip()
    return content.strip()


# ─── DOCX Builder ─────────────────────────────────────────────────────────────

def add_cover_page(doc, project_name: str, profile: str, doc_count: int):
    """Add an executive cover page."""
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading(project_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Business Document Corpus")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Profile: {profile}   |   Documents: {doc_count}   |   Generated: {datetime.now().strftime('%Y-%m-%d')}")
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()


def add_toc(doc, documents: dict):
    """Add a simple table of contents."""
    doc.add_heading("Table of Contents", level=1)
    for phase, doc_types in PHASE_ORDER.items():
        phase_docs = [dt for dt in doc_types if dt in documents]
        if phase_docs:
            doc.add_heading(phase.title() + " Layer", level=2)
            for dt in phase_docs:
                p = doc.add_paragraph(f"  {DOCUMENT_TITLES.get(dt, dt)}", style="List Bullet")
    doc.add_page_break()


def markdown_to_docx(doc, content: str, doc_type: str):
    """Convert markdown content to docx paragraphs (simplified)."""
    lines = content.split("\n")
    in_table = False
    table_rows = []

    for line in lines:
        # Headings
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        # Tables
        elif line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(set(c.strip()) <= set("-: ") for c in cells):
                # Separator row — skip
                continue
            table_rows.append(cells)
            in_table = True
        else:
            # Flush table
            if in_table and table_rows:
                max_cols = max(len(r) for r in table_rows)
                t = doc.add_table(rows=len(table_rows), cols=max_cols)
                t.style = "Table Grid"
                for i, row in enumerate(table_rows):
                    for j, cell_text in enumerate(row):
                        if j < max_cols:
                            t.rows[i].cells[j].text = cell_text
                            if i == 0:
                                for para in t.rows[i].cells[j].paragraphs:
                                    for run in para.runs:
                                        run.bold = True
                table_rows = []
                in_table = False

            # Code blocks
            if line.startswith("```"):
                continue
            # Bullet points
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            elif re.match(r"^\d+\. ", line):
                doc.add_paragraph(re.sub(r"^\d+\. ", "", line).strip(), style="List Number")
            # Horizontal rule
            elif line.strip() in ("---", "===", "***"):
                doc.add_paragraph("─" * 60)
            # Empty line
            elif line.strip() == "":
                doc.add_paragraph()
            # Normal paragraph
            else:
                # Strip inline markdown (bold, italic, code)
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
                clean = re.sub(r"\*(.+?)\*", r"\1", clean)
                clean = re.sub(r"`(.+?)`", r"\1", clean)
                doc.add_paragraph(clean)


def build_docx(documents: dict, project_name: str, profile: str, output_path: Path):
    """Build the full workbook .docx."""
    if not DOCX_AVAILABLE:
        print("python-docx not available — skipping .docx generation")
        return

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Cover
    add_cover_page(doc, project_name, profile, len(documents))

    # TOC
    add_toc(doc, documents)

    # Documents in phase order
    for phase, doc_types in PHASE_ORDER.items():
        phase_docs = [dt for dt in doc_types if dt in documents]
        if not phase_docs:
            continue

        # Phase divider
        doc.add_heading(f"{phase.title()} Layer", level=1)
        doc.add_page_break()

        for doc_type in phase_docs:
            path = documents[doc_type]
            content = strip_frontmatter(read_document(path))

            # Document title
            doc.add_heading(DOCUMENT_TITLES.get(doc_type, doc_type), level=1)
            markdown_to_docx(doc, content, doc_type)
            doc.add_page_break()

    doc.save(str(output_path))
    print(f"✓ Workbook saved: {output_path}")


# ─── ZIP Builder ──────────────────────────────────────────────────────────────

def build_zip(documents: dict, project_name: str, output_path: Path):
    """Build organized folder structure as .zip."""
    slug = re.sub(r"[^a-z0-9]+", "-", project_name.lower()).strip("-")
    tmp_dir = output_path.parent / f"_tmp_{slug}"
    tmp_dir.mkdir(parents=True, exist_ok=True)

    for phase, doc_types in PHASE_ORDER.items():
        for doc_type in doc_types:
            if doc_type not in documents:
                continue
            phase_dir = tmp_dir / f"{phase}-layer"
            phase_dir.mkdir(exist_ok=True)
            src = documents[doc_type]
            dst = phase_dir / f"{doc_type}.md"
            shutil.copy2(src, dst)

    # Add corpus summary if it exists
    summary_candidates = list(output_path.parent.glob("**/corpus_summary*.md"))
    if summary_candidates:
        shutil.copy2(summary_candidates[0], tmp_dir / "CORPUS_SUMMARY.md")

    # Zip
    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for file in tmp_dir.rglob("*"):
            zf.write(file, file.relative_to(tmp_dir.parent))

    shutil.rmtree(tmp_dir)
    print(f"✓ Folder structure saved: {output_path}")


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Package approved corpus into workbook + zip")
    parser.add_argument("--corpus-dir", required=True, help="Directory containing approved .md documents")
    parser.add_argument("--project-name", required=True, help="Project name for cover page")
    parser.add_argument("--profile", default="product", help="Workflow profile used")
    parser.add_argument("--output-dir", default="./output/final", help="Output directory")
    args = parser.parse_args()

    corpus_dir = Path(args.corpus_dir)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"\n📂 Scanning corpus: {corpus_dir}")
    documents = find_corpus_files(corpus_dir)

    if not documents:
        print("⚠ No documents found in corpus directory.")
        return

    print(f"✓ Found {len(documents)} documents: {', '.join(documents.keys())}")

    slug = re.sub(r"[^a-z0-9]+", "-", args.project_name.lower()).strip("-")
    date_str = datetime.now().strftime("%Y%m%d")

    # Build .docx
    docx_path = output_dir / f"{slug}-corpus-{date_str}.docx"
    build_docx(documents, args.project_name, args.profile, docx_path)

    # Build .zip
    zip_path = output_dir / f"{slug}-corpus-{date_str}.zip"
    build_zip(documents, args.project_name, zip_path)

    print(f"\n✅ Corpus packaged successfully.")
    print(f"   Workbook: {docx_path}")
    print(f"   Archive:  {zip_path}")


if __name__ == "__main__":
    main()
